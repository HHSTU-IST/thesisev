"""Parsing and preprocessing utilities."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast
from zipfile import BadZipFile, ZipFile, ZipInfo

from docx import Document as load_docx_document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn

from thesisev.models import Paragraph, Section, Sentence, ThesisDocument

CHINESE_NUMERALS = "一二三四五六七八九十百千万零"
SECTION_PATTERN = re.compile(
    r"^(?P<prefix>("
    r"第[" + CHINESE_NUMERALS + r"]+[章节部分]"
    r"|第[" + CHINESE_NUMERALS + r"]+节"
    r"|[0-9]+(?:\.[0-9]+){0,3}"
    r"|[IVXLC]+(?:\.[IVXLC]+){0,3}"
    r"|[(（]?[0-9一二三四五六七八九十]+[)）]"
    r"))[\s、.\-:：]+(?P<title>.+)$",
    flags=re.IGNORECASE,
)
MARKDOWN_HEADING_PATTERN = re.compile(r"^\s*#+\s*(.+?)\s*$")
SENTENCE_SPLIT_PATTERN = re.compile(r"(?<=[。！？!?；;])\s*")
MERMAID_FENCE_PATTERN = re.compile(r"(?ms)^```[ \t]*mermaid[^\n]*\n.*?^```[ \t]*$")
COVER_TITLE_PATTERNS = (
    re.compile(r"^\s*课题名称\s*[：:]\s*(?P<title>.+?)\s*$"),
    re.compile(r"^\s*(?:论文题目|题目)\s*[：:]\s*(?P<title>.+?)\s*$"),
)
MAX_DOCX_ARCHIVE_BYTES = 25 * 1024 * 1024
MAX_DOCX_MEMBER_COUNT = 512
MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES = 80 * 1024 * 1024
MAX_DOCX_MEMBER_BYTES = 30 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_DOCX_TEXT_PARAGRAPHS = 5000
MAX_DOCX_SNAPSHOT_PARAGRAPHS = 5000
MAX_DOCX_SNAPSHOT_TABLES = 500
MAX_DOCX_SNAPSHOT_SECTIONS = 200
MAX_DOCX_RUNS_PER_PARAGRAPH = 500


@dataclass(slots=True)
class HeadingMatch:
    """A detected heading line."""

    line_index: int
    level: int
    title: str
    numbering: str
    heading: str


def load_document(path: str | Path) -> ThesisDocument:
    """Load and parse a thesis from markdown or docx."""

    source_path = Path(path)
    if source_path.suffix.lower() == ".docx":
        docx_document = open_docx_document(source_path)
        raw_text = read_docx_text_from_document(docx_document)
        format_snapshot = read_docx_format_snapshot_from_document(docx_document)
        cleaned_text = clean_text(raw_text)
        front_matter, sections = parse_docx_sections(docx_document)
        title = extract_document_title(front_matter, fallback=source_path.stem)
    else:
        raw_text = read_source_text(source_path)
        format_snapshot: dict[str, Any] = {}
        cleaned_text = clean_text(raw_text)
        lines = [line.strip() for line in cleaned_text.splitlines()]
        title = next((line for line in lines if line), source_path.stem)
        body_text = remove_title_from_text(cleaned_text, title)
        front_matter, sections = parse_sections(body_text)
    root_sections = build_section_tree(sections)
    paragraphs = collect_document_paragraphs(front_matter, sections)
    sentences = flatten_sentence_text(paragraphs)
    total_word_count = sum(paragraph.word_count for paragraph in paragraphs)
    abstract = find_abstract(sections, front_matter)
    if isinstance(format_snapshot, dict):
        format_snapshot["word_count"] = total_word_count
        format_snapshot["section_count"] = len(sections)

    return ThesisDocument(
        title=title,
        source_path=str(source_path),
        source_type=source_path.suffix.lower().lstrip("."),
        raw_text=raw_text,
        cleaned_text=cleaned_text,
        front_matter=front_matter,
        abstract=abstract,
        sections=sections,
        root_sections=root_sections,
        paragraphs=paragraphs,
        sentences=sentences,
        total_word_count=total_word_count,
        format_snapshot=format_snapshot,
    )


def extract_document_title(text: str, *, fallback: str) -> str:
    """Extract a cover-page title, preferring explicit topic metadata."""

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for pattern in COVER_TITLE_PATTERNS:
        for line in lines:
            match = pattern.match(line)
            if match is None:
                continue
            candidate = cast(str, match.group("title")).strip()
            if is_document_title_candidate(candidate):
                return candidate
    return next((line for line in lines if is_document_title_candidate(line)), fallback)


def is_document_title_candidate(text: str) -> bool:
    """Reject common cover labels and template instructions as title candidates."""

    if not text or len(text) > 120:
        return False
    if re.match(
        r"^(课程名称|专业班级|课题名称|学生学号|学生姓名|所属院部|指导教师|"
        r"作业要求|要求|作业指南|文字格式|图表格式)\s*[：:]?",
        text,
    ):
        return False
    return not text.startswith(("居中", "黑体", "宋体", "颜色为", "不要", "使用"))


def read_source_text(path: Path) -> str:
    """Read thesis text from a supported source file."""

    suffix = path.suffix.lower()
    if suffix == ".md":
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return read_docx_text(path)
    msg = f"unsupported file type: {path.suffix or 'unknown'}"
    raise ValueError(msg)


def read_docx_text(path: Path) -> str:
    """Extract plain text from a docx file using python-docx."""

    return read_docx_text_from_document(open_docx_document(path))


def read_docx_text_from_document(document: Any) -> str:
    """Extract plain text from a python-docx document object."""

    paragraphs: list[str] = []
    for index, paragraph in enumerate(iter_docx_paragraphs(document), start=1):
        if index > MAX_DOCX_TEXT_PARAGRAPHS:
            raise_docx_traversal_limit("paragraphs", MAX_DOCX_TEXT_PARAGRAPHS)
        text = normalize_docx_paragraph_text(getattr(paragraph, "text", ""))
        if not text:
            continue
        if len(getattr(paragraph, "runs", [])) > MAX_DOCX_RUNS_PER_PARAGRAPH:
            raise_docx_traversal_limit(
                "runs per paragraph", MAX_DOCX_RUNS_PER_PARAGRAPH
            )
        paragraphs.append(text)
    return "\n\n".join(paragraphs)


def read_docx_format_snapshot(path: Path) -> dict[str, Any]:
    """Extract a compact formatting snapshot from a docx file."""

    return read_docx_format_snapshot_from_document(open_docx_document(path))


def read_docx_format_snapshot_from_document(document: Any) -> dict[str, Any]:
    """Extract a compact formatting snapshot from a python-docx document."""

    style_map = build_docx_style_map(document)
    paragraphs: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    sections: list[dict[str, Any]] = []

    for index, paragraph in enumerate(iter_docx_paragraphs(document), start=1):
        if index > MAX_DOCX_SNAPSHOT_PARAGRAPHS:
            raise_docx_traversal_limit(
                "paragraph snapshots", MAX_DOCX_SNAPSHOT_PARAGRAPHS
            )
        snapshot = extract_docx_paragraph_snapshot(paragraph, style_map=style_map)
        if snapshot["text"] or snapshot["style"] or snapshot["runs"]:
            paragraphs.append(snapshot)

    for index, table in enumerate(iter_docx_tables(document), start=1):
        if index > MAX_DOCX_SNAPSHOT_TABLES:
            raise_docx_traversal_limit("tables", MAX_DOCX_SNAPSHOT_TABLES)
        tables.append(extract_docx_table_snapshot(table))

    for index, section in enumerate(getattr(document, "sections", []), start=1):
        if index > MAX_DOCX_SNAPSHOT_SECTIONS:
            raise_docx_traversal_limit("sections", MAX_DOCX_SNAPSHOT_SECTIONS)
        sections.append(extract_docx_section_snapshot(section))

    if not sections:
        sections.append({})

    return {"paragraphs": paragraphs, "tables": tables, "sections": sections}


def parse_docx_sections(document: Any) -> tuple[str, list[Section]]:
    """Parse section structure from a python-docx document."""

    lines: list[str] = []
    headings: list[HeadingMatch] = []
    paragraphs = list(getattr(document, "paragraphs", []))
    has_styled_headings = any(
        match_docx_style_heading(
            paragraph, normalize_docx_paragraph_text(paragraph.text)
        )
        is not None
        for paragraph in paragraphs
    )

    for paragraph in paragraphs:
        paragraph_text = normalize_docx_paragraph_text(paragraph.text)
        if not paragraph_text:
            continue
        lines.append(paragraph_text)
        heading = match_docx_paragraph_heading(
            paragraph, paragraph_text, allow_text_fallback=not has_styled_headings
        )
        if heading is not None:
            headings.append(
                HeadingMatch(
                    line_index=len(lines) - 1,
                    level=heading[0],
                    title=heading[1],
                    numbering=heading[2],
                    heading=heading[3],
                )
            )

    if not headings:
        msg = "docx file must contain explicit heading markers"
        raise ValueError(msg)

    first_level_one = next(
        (index for index, heading in enumerate(headings) if heading.level == 1), None
    )
    if first_level_one is not None:
        headings = headings[first_level_one:]

    front_matter = "\n".join(lines[: headings[0].line_index]).strip()
    sections: list[Section] = []
    for position, heading in enumerate(headings):
        start_index = heading.line_index + 1
        end_index = len(lines)
        if position + 1 < len(headings):
            end_index = headings[position + 1].line_index
        content = "\n".join(lines[start_index:end_index]).strip()
        sections.append(
            build_section(
                identifier=str(position + 1),
                level=heading.level,
                title=heading.title,
                heading=heading.heading,
                numbering=heading.numbering,
                content=content,
            )
        )
    return front_matter, sections


def open_docx_document(path: Path) -> Any:
    """Open a DOCX file through python-docx after lightweight archive checks."""

    validate_docx_archive_size(path)
    try:
        with ZipFile(path) as archive:
            validate_docx_archive_members(archive)
            validate_docx_member_size(
                get_docx_document_xml_info(archive), max_size=MAX_DOCX_MEMBER_BYTES
            )
            if "word/styles.xml" in archive.namelist():
                validate_docx_member_size(
                    archive.getinfo("word/styles.xml"), max_size=MAX_DOCX_MEMBER_BYTES
                )
    except BadZipFile as exc:
        msg = "invalid docx archive"
        raise ValueError(msg) from exc
    try:
        return load_docx_document(str(path))
    except (BadZipFile, PackageNotFoundError) as exc:
        msg = "invalid docx document"
        raise ValueError(msg) from exc


def iter_docx_paragraphs(document: Any):
    """Iterate over document paragraphs in reading order."""

    if hasattr(document, "paragraphs"):
        yield from document.paragraphs
    if hasattr(document, "tables"):
        for table in document.tables:
            yield from iter_docx_table_paragraphs(table)


def iter_docx_table_paragraphs(table: Any):
    """Iterate paragraphs from a table and its nested cells."""

    for row in getattr(table, "rows", []):
        for cell in getattr(row, "cells", []):
            yield from getattr(cell, "paragraphs", [])
            for nested_table in getattr(cell, "tables", []):
                yield from iter_docx_table_paragraphs(nested_table)


def iter_docx_tables(document: Any):
    """Iterate over all document tables."""

    if hasattr(document, "tables"):
        for table in document.tables:
            yield table
            yield from iter_docx_nested_tables(table)


def iter_docx_nested_tables(table: Any):
    """Iterate nested tables from a table."""

    for row in getattr(table, "rows", []):
        for cell in getattr(row, "cells", []):
            for nested_table in getattr(cell, "tables", []):
                yield nested_table
                yield from iter_docx_nested_tables(nested_table)


def iter_docx_runs(paragraph: Any):
    """Iterate runs from a paragraph."""

    yield from getattr(paragraph, "runs", [])


def normalize_docx_paragraph_text(text: str) -> str:
    """Normalize python-docx paragraph text for section parsing."""

    normalized = (
        (text or "").replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
    )
    normalized = normalized.replace("\t", " ")
    normalized = re.sub(r"[ ]+", " ", normalized)
    return normalized.strip()


def build_docx_style_map(document: Any) -> dict[str, dict[str, Any]]:
    """Build a style snapshot map from a python-docx document."""

    style_map: dict[str, dict[str, Any]] = {}
    for style in getattr(document, "styles", []):
        style_id = getattr(style, "style_id", None)
        if not style_id:
            continue
        base_style = getattr(style, "base_style", None)
        style_map[str(style_id)] = {
            "name": normalize_docx_style_name(getattr(style, "name", None))
            or getattr(style, "name", None),
            "based_on": getattr(base_style, "style_id", None),
            "font_name": read_docx_font_name(style),
            "font_size_pt": parse_docx_length(
                getattr(getattr(style, "font", None), "size", None)
            ),
            "bold": getattr(getattr(style, "font", None), "bold", None),
            "alignment": normalize_docx_alignment(
                getattr(
                    getattr(style.paragraph_format, "alignment", None), "name", None
                )
            )
            if getattr(style, "paragraph_format", None) is not None
            else None,
            "line_spacing": parse_docx_line_spacing_from_style(style),
            "space_before_pt": parse_docx_length(
                getattr(getattr(style, "paragraph_format", None), "space_before", None)
            ),
            "space_after_pt": parse_docx_length(
                getattr(getattr(style, "paragraph_format", None), "space_after", None)
            ),
            "first_line_indent_pt": parse_docx_length(
                getattr(
                    getattr(style, "paragraph_format", None), "first_line_indent", None
                )
            ),
        }
    return resolve_docx_style_map(style_map)


def resolve_docx_style_map(
    style_map: dict[str, dict[str, Any]],
) -> dict[str, dict[str, Any]]:
    """Resolve inherited style properties in a style map."""

    resolved: dict[str, dict[str, Any]] = {}

    def resolve_style(style_id: str) -> dict[str, Any]:
        if style_id in resolved:
            return resolved[style_id]
        style = style_map.get(style_id, {})
        base_style_id = style.get("based_on")
        base: dict[str, Any] = resolve_style(base_style_id) if base_style_id else {}
        merged = {
            "name": style.get("name") or base.get("name"),
            "based_on": base_style_id,
            "font_name": style.get("font_name") or base.get("font_name"),
            "font_size_pt": style.get("font_size_pt") or base.get("font_size_pt"),
            "bold": style.get("bold")
            if style.get("bold") is not None
            else base.get("bold"),
            "alignment": style.get("alignment") or base.get("alignment"),
            "line_spacing": style.get("line_spacing") or base.get("line_spacing"),
            "space_before_pt": style.get("space_before_pt")
            if style.get("space_before_pt") is not None
            else base.get("space_before_pt"),
            "space_after_pt": style.get("space_after_pt")
            if style.get("space_after_pt") is not None
            else base.get("space_after_pt"),
            "first_line_indent_pt": style.get("first_line_indent_pt")
            if style.get("first_line_indent_pt") is not None
            else base.get("first_line_indent_pt"),
        }
        resolved[style_id] = merged
        return merged

    for style_id in style_map:
        resolve_style(style_id)
    return resolved


def validate_docx_archive_size(path: Path) -> None:
    """Reject archive files that are too large before opening them."""

    if path.stat().st_size > MAX_DOCX_ARCHIVE_BYTES:
        msg = (
            "docx archive is too large; "
            f"maximum size is {MAX_DOCX_ARCHIVE_BYTES // (1024 * 1024)} MiB"
        )
        raise ValueError(msg)


def validate_docx_archive_members(archive: ZipFile) -> None:
    """Validate zip metadata before any member is decompressed."""

    members = archive.infolist()
    if len(members) > MAX_DOCX_MEMBER_COUNT:
        msg = (
            f"docx archive has too many zip members; maximum is {MAX_DOCX_MEMBER_COUNT}"
        )
        raise ValueError(msg)

    total_uncompressed = 0
    for member in members:
        if member.is_dir():
            continue
        validate_docx_member_size(member, max_size=MAX_DOCX_MEMBER_BYTES)
        total_uncompressed += member.file_size
        if total_uncompressed > MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES:
            msg = (
                "docx archive expands to too much data; "
                f"maximum is {MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES // (1024 * 1024)} MiB"
            )
            raise ValueError(msg)


def get_docx_document_xml_info(archive: ZipFile) -> ZipInfo:
    """Return the central-directory entry for word/document.xml."""

    try:
        return archive.getinfo("word/document.xml")
    except KeyError as exc:
        msg = "docx archive is missing word/document.xml"
        raise ValueError(msg) from exc


def raise_docx_traversal_limit(label: str, max_items: int) -> None:
    """Raise a consistent DOCX traversal-limit error."""

    msg = f"docx document has too many {label}; maximum is {max_items}"
    raise ValueError(msg)


def validate_docx_member_size(member: ZipInfo, *, max_size: int) -> None:
    """Reject suspicious or oversized zip members before decompression."""

    if member.file_size > max_size:
        msg = (
            f"docx member {member.filename} is too large; "
            f"maximum size is {max_size // (1024 * 1024)} MiB"
        )
        raise ValueError(msg)
    if member.file_size and member.compress_size == 0:
        msg = f"docx member {member.filename} has suspicious compression metadata"
        raise ValueError(msg)
    if member.compress_size:
        ratio = member.file_size / member.compress_size
        if ratio > MAX_DOCX_COMPRESSION_RATIO:
            msg = f"docx member {member.filename} has suspicious compression ratio"
            raise ValueError(msg)


def extract_docx_paragraph_snapshot(
    paragraph: Any, *, style_map: dict[str, dict[str, Any]] | None = None
) -> dict[str, Any]:
    """Extract a paragraph-level formatting snapshot."""

    paragraph_style_id = get_docx_paragraph_style_id(paragraph)
    style_snapshot = resolve_docx_style_snapshot(paragraph_style_id, style_map)
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    runs: list[dict[str, Any]] = []
    for index, run in enumerate(getattr(paragraph, "runs", []), start=1):
        if index > MAX_DOCX_RUNS_PER_PARAGRAPH:
            raise_docx_traversal_limit(
                "runs per paragraph", MAX_DOCX_RUNS_PER_PARAGRAPH
            )
        run_snapshot = extract_docx_run_snapshot(
            run, style_map=style_map, paragraph_style_snapshot=style_snapshot
        )
        if run_snapshot["text"] or any(
            run_snapshot[key] is not None
            for key in ("font_name", "font_size_pt", "bold")
        ):
            runs.append(run_snapshot)

    text = normalize_docx_paragraph_text(getattr(paragraph, "text", ""))
    primary_run: dict[str, Any] = next(
        (run for run in runs if run.get("text")), runs[0] if runs else {}
    )
    return {
        "text": text,
        "style": style_snapshot.get("name"),
        "style_id": paragraph_style_id,
        "alignment": paragraph_alignment_to_token(
            getattr(paragraph_format, "alignment", None)
        )
        or style_snapshot.get("alignment"),
        "line_spacing": parse_docx_line_spacing_from_paragraph(paragraph)
        or style_snapshot.get("line_spacing"),
        "space_before_pt": parse_docx_length(
            getattr(paragraph_format, "space_before", None)
        )
        if paragraph_format is not None and paragraph_format.space_before is not None
        else style_snapshot.get("space_before_pt"),
        "space_after_pt": parse_docx_length(
            getattr(paragraph_format, "space_after", None)
        )
        if paragraph_format is not None and paragraph_format.space_after is not None
        else style_snapshot.get("space_after_pt"),
        "first_line_indent_pt": parse_docx_length(
            getattr(paragraph_format, "first_line_indent", None)
        )
        if paragraph_format is not None
        and paragraph_format.first_line_indent is not None
        else style_snapshot.get("first_line_indent_pt"),
        "runs": runs,
        "run": primary_run,
    }


def extract_docx_run_snapshot(
    run: Any,
    *,
    style_map: dict[str, dict[str, Any]] | None = None,
    paragraph_style_snapshot: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Extract a run-level formatting snapshot."""

    run_style_snapshot = resolve_docx_run_style_snapshot(run, style_map)
    paragraph_style_snapshot = paragraph_style_snapshot or {}
    return {
        "text": normalize_docx_paragraph_text(getattr(run, "text", "")),
        "font_name": read_docx_font_name(run)
        or run_style_snapshot.get("font_name")
        or paragraph_style_snapshot.get("font_name"),
        "font_size_pt": parse_docx_length(
            getattr(getattr(run, "font", None), "size", None)
        )
        or run_style_snapshot.get("font_size_pt")
        or paragraph_style_snapshot.get("font_size_pt"),
        "bold": getattr(getattr(run, "font", None), "bold", None),
    }


def read_docx_font_name(element_owner: Any) -> str | None:
    """Read East Asian fonts before western fonts from a python-docx object."""

    element = getattr(element_owner, "_element", None)
    properties = getattr(element, "rPr", None)
    fonts = getattr(properties, "rFonts", None)
    if fonts is not None:
        for attribute in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            value = fonts.get(qn(attribute))
            if value:
                return str(value)
    font = getattr(element_owner, "font", None)
    name = getattr(font, "name", None)
    return name if isinstance(name, str) else None


def extract_docx_table_snapshot(table: Any) -> dict[str, Any]:
    """Extract a table-level formatting snapshot."""

    return {
        "alignment": paragraph_alignment_to_token(getattr(table, "alignment", None))
    }


def extract_docx_section_snapshot(section: Any) -> dict[str, Any]:
    """Extract a section-level formatting snapshot."""

    return {
        "top_margin_pt": parse_docx_length(getattr(section, "top_margin", None)),
        "bottom_margin_pt": parse_docx_length(getattr(section, "bottom_margin", None)),
        "left_margin_pt": parse_docx_length(getattr(section, "left_margin", None)),
        "right_margin_pt": parse_docx_length(getattr(section, "right_margin", None)),
    }


def get_docx_paragraph_style_id(paragraph: Any) -> str | None:
    """Return a paragraph style id from python-docx objects."""

    style = getattr(paragraph, "style", None)
    if style is None:
        return None
    style_id = getattr(style, "style_id", None)
    if style_id:
        return str(style_id)
    name = getattr(style, "name", None)
    if not name:
        return None
    return normalize_docx_style_name(str(name))


def parse_docx_length(value: Any) -> float | None:
    """Convert a python-docx length object to points."""

    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except AttributeError:
        try:
            return round(float(value) / 12700, 2)
        except (TypeError, ValueError):
            return None


def resolve_docx_run_style_snapshot(
    run: Any, style_map: dict[str, dict[str, Any]] | None = None
) -> dict[str, Any]:
    """Resolve run style inheritance if direct properties are missing."""

    style = getattr(run, "style", None)
    style_id = getattr(style, "style_id", None) if style is not None else None
    if not style_id and style is not None:
        style_id = getattr(style, "name", None)
    if not style_id:
        return {}
    return resolve_docx_style_snapshot(str(style_id), style_map)


def parse_docx_line_spacing_from_paragraph(paragraph: Any) -> float | None:
    """Read line spacing from a python-docx paragraph."""

    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is None:
        return None
    line_spacing = getattr(paragraph_format, "line_spacing", None)
    if line_spacing is None:
        return None
    if isinstance(line_spacing, (int, float)):
        return round(float(line_spacing), 2)
    try:
        return round(float(line_spacing), 2)
    except (TypeError, ValueError):
        return parse_docx_length(line_spacing)


def parse_docx_line_spacing_from_style(style: Any) -> float | None:
    """Read line spacing from a python-docx style object."""

    paragraph_format = getattr(style, "paragraph_format", None)
    if paragraph_format is None:
        return None
    line_spacing = getattr(paragraph_format, "line_spacing", None)
    if line_spacing is None:
        return None
    if isinstance(line_spacing, (int, float)):
        return round(float(line_spacing), 2)
    return parse_docx_length(line_spacing)


def normalize_docx_style_name(value: str | None) -> str | None:
    """Normalize a docx style identifier into a display name."""

    if not value:
        return None
    normalized = value.replace("_", " ").strip()
    normalized = re.sub(r"\s+", " ", normalized)
    heading_match = re.match(r"(?i)^heading\s*(\d+)$", normalized)
    if heading_match is not None:
        return f"Heading {heading_match.group(1)}"
    if re.match(r"(?i)^normal(?:\s*\(web\))?$", normalized):
        return "Normal"
    return re.sub(r"(?<=\D)(\d+)$", r" \1", normalized)


def resolve_docx_style_snapshot(
    style_id: str | None, style_map: dict[str, dict[str, Any]] | None = None
) -> dict[str, Any]:
    """Resolve a paragraph style id to its effective formatting snapshot."""

    if not style_id:
        return {}
    if not style_map:
        return {"name": normalize_docx_style_name(style_id)}
    style = style_map.get(style_id)
    if style is None:
        return {"name": normalize_docx_style_name(style_id)}
    base = resolve_docx_style_snapshot(style.get("based_on"), style_map)
    return {
        "name": normalize_docx_style_name(style.get("name")) or base.get("name"),
        "font_name": style.get("font_name") or base.get("font_name"),
        "font_size_pt": style.get("font_size_pt") or base.get("font_size_pt"),
        "bold": style.get("bold")
        if style.get("bold") is not None
        else base.get("bold"),
        "line_spacing": style.get("line_spacing") or base.get("line_spacing"),
        "space_before_pt": style.get("space_before_pt")
        if style.get("space_before_pt") is not None
        else base.get("space_before_pt"),
        "space_after_pt": style.get("space_after_pt")
        if style.get("space_after_pt") is not None
        else base.get("space_after_pt"),
        "first_line_indent_pt": style.get("first_line_indent_pt")
        if style.get("first_line_indent_pt") is not None
        else base.get("first_line_indent_pt"),
    }


def normalize_docx_alignment(value: Any) -> str | None:
    """Normalize alignment values from python-docx."""

    if value is None:
        return None
    token_value = getattr(value, "name", None)
    token = str(value) if token_value is None else cast(str, token_value)
    token = token.replace("WD_PARAGRAPH_ALIGNMENT.", "").replace("_", " ")
    token = token.strip().lower()
    aliases = {
        "both": "justify",
        "justified": "justify",
        "center": "center",
        "centre": "center",
        "left": "left",
        "right": "right",
        "distribute": "distribute",
    }
    return aliases.get(token, token)


def paragraph_alignment_to_token(value: Any) -> str | None:
    """Convert a python-docx alignment value to a normalized token."""

    return normalize_docx_alignment(value)


def clean_text(text: str) -> str:
    """Normalize spacing and drop obvious markdown noise."""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = re.sub(r"^\s*#+\s*", "", normalized, flags=re.MULTILINE)
    normalized = re.sub(r"^\s*[-*_]{3,}\s*$", "", normalized, flags=re.MULTILINE)
    normalized = re.sub(r"^\s*>\s?", "", normalized, flags=re.MULTILINE)
    normalized = re.sub(r"^\s*\|\s*", "", normalized, flags=re.MULTILINE)
    return normalized.strip()


def remove_title_from_text(text: str, title: str) -> str:
    """Remove the title line from the text body if present."""

    lines = text.splitlines()
    skipped = False
    kept_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not skipped and stripped == title:
            skipped = True
            continue
        kept_lines.append(line)
    return "\n".join(kept_lines).strip()


def parse_sections(text: str) -> tuple[str, list[Section]]:
    """Split text into thesis sections using common heading patterns."""

    lines = text.splitlines()
    headings = [
        HeadingMatch(
            line_index=index,
            level=heading[0],
            title=heading[1],
            numbering=heading[2],
            heading=heading[3],
        )
        for index, raw_line in enumerate(lines)
        if (stripped := raw_line.strip())
        and (heading := match_section_heading(stripped)) is not None
    ]

    if not headings:
        msg = "markdown file must contain explicit heading markers"
        raise ValueError(msg)

    first_heading_line = headings[0].line_index
    front_matter = "\n".join(lines[:first_heading_line]).strip()
    sections: list[Section] = []
    for position, heading in enumerate(headings):
        end_index = len(lines)
        if position + 1 < len(headings):
            end_index = headings[position + 1].line_index
        content_lines = lines[heading.line_index + 1 : end_index]
        content = "\n".join(content_lines).strip()
        sections.append(
            build_section(
                identifier=str(position + 1),
                level=heading.level,
                title=heading.title,
                heading=heading.heading,
                numbering=heading.numbering,
                content=content,
            )
        )
    return front_matter, sections


def match_section_heading(line: str) -> tuple[int, str, str, str] | None:
    """Match a line against supported section heading formats."""

    markdown_match = MARKDOWN_HEADING_PATTERN.match(line)
    if markdown_match is not None:
        line = cast(str, markdown_match.group(1)).strip()

    if is_reference_heading(line):
        return 1, line, line, line

    match = SECTION_PATTERN.match(line)
    if match is not None:
        numbering = cast(str, match.group("prefix")).strip()
        title = cast(str, match.group("title")).strip()
        return infer_level(numbering), title, numbering, line
    return None


def match_docx_paragraph_heading(
    paragraph: Any, line: str, *, allow_text_fallback: bool = True
) -> tuple[int, str, str, str] | None:
    """Match a DOCX paragraph as a heading using style and text heuristics."""

    styled_heading = match_docx_style_heading(paragraph, line)
    if styled_heading is not None:
        return styled_heading
    if allow_text_fallback:
        return match_section_heading(line)
    return None


def match_docx_style_heading(
    paragraph: Any, line: str
) -> tuple[int, str, str, str] | None:
    """Match a DOCX paragraph as a heading using its Word style."""

    style = getattr(paragraph, "style", None)
    style_name = normalize_docx_style_name(getattr(style, "name", None))
    if style_name and style_name.lower().startswith("heading"):
        level_match = re.search(r"(\d+)", style_name)
        level = int(level_match.group(1)) if level_match else 1
        title = line.strip()
        return level, title, title, line
    return None


def infer_level(prefix: str) -> int:
    """Infer a section depth from its numbering prefix."""

    normalized = prefix.strip()
    if normalized.startswith("第") and any(
        token in normalized for token in ("章", "部分")
    ):
        return 1
    if normalized.startswith("第") and "节" in normalized:
        return 2
    if normalized.startswith(("(", "（")) and normalized.endswith((")", "）")):
        return 3
    if "." in normalized:
        return normalized.count(".") + 1
    if normalized.isdigit():
        return 1
    return 1


def build_section(
    identifier: str, level: int, title: str, heading: str, numbering: str, content: str
) -> Section:
    """Construct a section model from raw content."""

    paragraphs = build_paragraphs(content)
    is_mermaid_code = is_mermaid_heading(level=level, title=title)
    skip_format_check = is_mermaid_code or is_reference_heading(title)
    if skip_format_check:
        paragraphs = [
            mark_paragraph_skip_format_check(paragraph, is_mermaid_code=is_mermaid_code)
            for paragraph in paragraphs
        ]
    sentences = flatten_sentence_text(paragraphs)
    word_count = count_words(content)
    return Section(
        identifier=identifier,
        level=level,
        title=title,
        heading=heading,
        numbering=numbering,
        content=content,
        paragraphs=paragraphs,
        sentences=sentences,
        word_count=word_count,
        is_mermaid_code=is_mermaid_code,
        skip_format_check=skip_format_check,
    )


def is_mermaid_heading(*, level: int, title: str) -> bool:
    """Return whether a level-2/3 heading is labeled as Mermaid code in Chinese."""

    return level in {2, 3} and "Mermaid代码" in title


def is_reference_heading(title: str) -> bool:
    """Return whether a heading contains the Chinese references label."""

    return "参考文献" in title


def mark_paragraph_skip_format_check(
    paragraph: Paragraph, *, is_mermaid_code: bool
) -> Paragraph:
    """Return a paragraph copy marked to skip prose format checks."""

    return Paragraph(
        index=paragraph.index,
        text=paragraph.text,
        sentences=[],
        word_count=paragraph.word_count,
        is_mermaid_code=is_mermaid_code,
        skip_format_check=True,
        topic_relevance_score=paragraph.topic_relevance_score,
        topic_matched_keywords=paragraph.topic_matched_keywords,
        topic_is_relevant=paragraph.topic_is_relevant,
    )


def build_section_tree(sections: list[Section]) -> list[Section]:
    """Build a hierarchy of sections from a flat sequence."""

    root_sections: list[Section] = []
    stack: list[Section] = []
    counters: dict[int, int] = {}
    for section in sections:
        section.children = []
        counters[section.level] = counters.get(section.level, 0) + 1
        for deeper_level in list(counters):
            if deeper_level > section.level:
                counters.pop(deeper_level)
        section.identifier = ".".join(
            str(counters[level]) for level in sorted(counters) if level <= section.level
        )
        while stack and stack[-1].level >= section.level:
            stack.pop()
        if stack:
            stack[-1].children.append(section)
        else:
            root_sections.append(section)
        stack.append(section)
    return root_sections


def collect_document_paragraphs(
    front_matter: str, sections: list[Section]
) -> list[Paragraph]:
    """Collect document paragraphs excluding heading lines."""

    paragraph_texts: list[str] = []
    if front_matter.strip():
        paragraph_texts.extend(split_paragraphs(front_matter))
    for section in sections:
        paragraph_texts.extend(paragraph.text for paragraph in section.paragraphs)
    combined_text = "\n\n".join(paragraph_texts)
    return build_paragraphs(combined_text)


def build_paragraphs(text: str) -> list[Paragraph]:
    """Split text into non-empty paragraphs with sentence objects."""

    paragraphs: list[Paragraph] = []
    parts = split_paragraph_parts(text)
    for index, (part, is_mermaid_code) in enumerate(parts, start=1):
        paragraph_text = part.strip()
        if not paragraph_text:
            continue
        sentence_objects = (
            [] if is_mermaid_code else build_sentence_objects(paragraph_text)
        )
        paragraphs.append(
            Paragraph(
                index=index,
                text=paragraph_text,
                sentences=sentence_objects,
                word_count=count_words(paragraph_text),
                is_mermaid_code=is_mermaid_code,
            )
        )
    return paragraphs


def split_paragraph_parts(text: str) -> list[tuple[str, bool]]:
    """Split text into paragraphs while preserving Mermaid fenced code blocks."""

    parts: list[tuple[str, bool]] = []
    cursor = 0
    for match in MERMAID_FENCE_PATTERN.finditer(text):
        parts.extend(
            (part, False)
            for part in split_plain_paragraphs(text[cursor : match.start()])
        )
        parts.append((match.group(0), True))
        cursor = match.end()
    parts.extend((part, False) for part in split_plain_paragraphs(text[cursor:]))
    return parts


def split_plain_paragraphs(text: str) -> list[str]:
    """Split non-code text into paragraph strings."""

    return [part for part in re.split(r"\n\s*\n", text) if part.strip()]


def build_sentence_objects(text: str) -> list[Sentence]:
    """Split text into sentence objects."""

    sentence_parts = [
        part.strip() for part in SENTENCE_SPLIT_PATTERN.split(text) if part.strip()
    ]
    return [
        Sentence(index=index, text=part)
        for index, part in enumerate(sentence_parts, start=1)
    ]


def flatten_sentence_text(paragraphs: list[Paragraph]) -> list[str]:
    """Flatten sentence objects to plain strings for compatibility."""

    return [
        sentence.text for paragraph in paragraphs for sentence in paragraph.sentences
    ]


def split_paragraphs(text: str) -> list[str]:
    """Split text into non-empty paragraph strings."""

    return [paragraph.text for paragraph in build_paragraphs(text)]


def split_sentences(text: str) -> list[str]:
    """Split text into non-empty sentence strings."""

    return flatten_sentence_text(build_paragraphs(text))


def find_abstract(sections: list[Section], front_matter: str) -> str:
    """Extract abstract-like content from the document."""

    for section in sections:
        normalized = section.title.strip().lower()
        if normalized in {"摘要", "abstract"}:
            return section.content
    return front_matter


def count_words(text: str) -> int:
    """Count Chinese characters and Latin tokens as a rough word count."""

    chinese_chars = re.findall(r"[\u4e00-\u9fff]", text)
    latin_tokens = re.findall(r"[A-Za-z0-9_]+", text)
    return len(chinese_chars) + len(latin_tokens)
